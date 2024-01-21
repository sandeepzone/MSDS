import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from configuration.msds_config import MSDSConfig
config: MSDSConfig = MSDSConfig()
config.load("configuration/msds_config.yml")

def save_doc(res, dest_path):
    # JSON data as provided
    json_data = res
    # Load JSON data
    data = json.loads(json_data)

    # Create a Word document
    doc = Document()

    title = doc.add_heading(config.MSDS_DOC_HEADING, level=0)
    
    paragraph_format = title.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0 
    # Function to add a row to the table
    def add_row(table, col_data):
        row_cells = table.add_row().cells
        for idx, val in enumerate(col_data):
            row_cells[idx].text = val
            
    #this function is new
    def add_section_row(table, col_data):
        row_cells = table.add_row().cells
        for idx, val in enumerate(col_data):
            paragraph = row_cells[idx].paragraphs[0]
            run = paragraph.add_run(val)
            run.bold = True


    # Identify all unique sections and properties
    sections = set()
    for chem in data["Chemicals"].values():
        for section in chem:
            sections.add(section)

    # Create a table
    chemicals = list(data["Chemicals"].keys())
    num_columns = len(chemicals) + 1  # Additional column for properties
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    # header_cells[0].text = 'Property / Chemical'
    for idx, chem in enumerate(chemicals, start=1):
        header_cells[idx].text = chem

    # Add rows for each section and property
    map_dict = config.MAP_DICT
    i=0
#     sub_headers = ["a)","b)","c)","d)","e)","f)"]
    for section in sections:
        section = map_dict.get(i)
        # Add a row for the section title
        add_section_row(table, [section] + [''] * (num_columns - 1))
        #add_row(table, [section] + [''] * (num_columns - 1)) -- old and working

        # Determine the properties in this section
        properties = set()
        for chem in data["Chemicals"].values():
            if section in chem and isinstance(chem,dict):
                if isinstance(chem[section],dict):
                    properties.update(chem[section].keys())
                elif isinstance(chem[section],str):
                    properties.update({""})
        #Add rows for each property in the section
        for prop in sorted(properties):  
            row_data = [prop]
            for chem in chemicals:
                if isinstance(data["Chemicals"][chem].get(section, {}),dict):
                    value = data["Chemicals"][chem].get(section, {}).get(prop, "")
                elif isinstance(data["Chemicals"][chem].get(section, {}),str):
                    value = data["Chemicals"][chem].get(section, {})
                row_data.append(value)
            add_row(table, row_data)
        i = i+1

    # Save the Word document
    doc.save(dest_path) 